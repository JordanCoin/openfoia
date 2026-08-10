"""Document ingestion from various sources."""

from __future__ import annotations

import asyncio
import hashlib
import mimetypes
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from uuid import uuid4

from ..models import DocumentType
from ..models import utcnow as _utcnow


@dataclass
class IngestResult:
    """Result of document ingestion."""

    document_id: str
    filename: str
    file_path: str
    file_size: int
    mime_type: str
    page_count: int | None
    extracted_text: str | None
    checksum: str
    metadata: dict[str, Any]


class DocumentIngester:
    """Ingest documents from various sources into the system."""

    def __init__(
        self,
        storage_path: Path | str,
        max_file_size_mb: int = 100,
    ):
        self.storage_path = Path(storage_path)
        self.storage_path.mkdir(parents=True, exist_ok=True)
        self.max_file_size = max_file_size_mb * 1024 * 1024

    async def ingest_file(
        self,
        file_path: Path | str,
        doc_type: DocumentType = DocumentType.FULL_RESPONSE,
        request_id: str | None = None,
        metadata: dict[str, Any] | None = None,
        strip_metadata: bool = True,
    ) -> IngestResult:
        """Ingest a file from the filesystem."""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_size = file_path.stat().st_size
        if file_size > self.max_file_size:
            raise ValueError(f"File too large: {file_size} bytes (max {self.max_file_size})")

        # Determine MIME type
        mime_type, _ = mimetypes.guess_type(str(file_path))
        mime_type = mime_type or "application/octet-stream"

        # Generate document ID
        doc_id = str(uuid4())

        # Calculate checksum
        checksum = await self._calculate_checksum(file_path)

        # Copy to storage
        dest_dir = self.storage_path / doc_id[:2] / doc_id[2:4]
        dest_dir.mkdir(parents=True, exist_ok=True)
        dest_path = dest_dir / f"{doc_id}{file_path.suffix}"

        await asyncio.to_thread(shutil.copy2, file_path, dest_path)

        # Strip metadata from the stored copy
        stripped_info: dict[str, Any] = {}
        if strip_metadata:
            from .metadata import strip_metadata as _strip_metadata

            stripped_info = await asyncio.to_thread(_strip_metadata, dest_path)

        # Get page count for PDFs
        page_count = None
        if mime_type == "application/pdf":
            page_count = await self._get_pdf_page_count(dest_path)

        # Auto-extract text based on file type
        extracted_text = await self._extract_text(dest_path, mime_type)

        return IngestResult(
            document_id=doc_id,
            filename=file_path.name,
            file_path=str(dest_path),
            file_size=file_size,
            mime_type=mime_type,
            page_count=page_count,
            extracted_text=extracted_text,
            checksum=checksum,
            metadata={
                "original_path": str(file_path),
                "ingested_at": _utcnow().isoformat(),
                "doc_type": doc_type.value,
                "request_id": request_id,
                "metadata_stripped": stripped_info if stripped_info else None,
                **(metadata or {}),
            },
        )

    async def ingest_bytes(
        self,
        content: bytes,
        filename: str,
        doc_type: DocumentType = DocumentType.FULL_RESPONSE,
        request_id: str | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> IngestResult:
        """Ingest raw bytes (e.g., from email attachment)."""
        if len(content) > self.max_file_size:
            raise ValueError(f"Content too large: {len(content)} bytes")

        # Determine MIME type
        mime_type, _ = mimetypes.guess_type(filename)
        mime_type = mime_type or "application/octet-stream"

        # Generate document ID
        doc_id = str(uuid4())

        # Calculate checksum
        checksum = hashlib.sha256(content).hexdigest()

        # Save to storage
        suffix = Path(filename).suffix
        dest_dir = self.storage_path / doc_id[:2] / doc_id[2:4]
        dest_dir.mkdir(parents=True, exist_ok=True)
        dest_path = dest_dir / f"{doc_id}{suffix}"

        await asyncio.to_thread(dest_path.write_bytes, content)

        # Get page count for PDFs
        page_count = None
        if mime_type == "application/pdf":
            page_count = await self._get_pdf_page_count(dest_path)

        # Auto-extract text
        extracted_text = await self._extract_text(dest_path, mime_type)

        return IngestResult(
            document_id=doc_id,
            filename=filename,
            file_path=str(dest_path),
            file_size=len(content),
            mime_type=mime_type,
            page_count=page_count,
            extracted_text=extracted_text,
            checksum=checksum,
            metadata={
                "ingested_at": _utcnow().isoformat(),
                "doc_type": doc_type.value,
                "request_id": request_id,
                **(metadata or {}),
            },
        )

    async def ingest_email_attachment(
        self,
        email_message: Any,  # email.message.Message
        attachment_index: int,
        request_id: str | None = None,
    ) -> IngestResult:
        """Extract and ingest an attachment from an email."""

        # Find attachment
        attachments = []
        for part in email_message.walk():
            if part.get_content_maintype() == "multipart":
                continue
            if part.get("Content-Disposition") is None:
                continue
            attachments.append(part)

        if attachment_index >= len(attachments):
            raise IndexError(f"Attachment index {attachment_index} out of range")

        attachment = attachments[attachment_index]
        filename = attachment.get_filename() or f"attachment_{attachment_index}"
        content = attachment.get_payload(decode=True)

        return await self.ingest_bytes(
            content=content,
            filename=filename,
            request_id=request_id,
            metadata={
                "source": "email",
                "email_subject": email_message.get("Subject"),
                "email_from": email_message.get("From"),
                "email_date": email_message.get("Date"),
            },
        )

    async def ingest_directory(
        self,
        dir_path: Path | str,
        recursive: bool = True,
        file_patterns: list[str] | None = None,
        request_id: str | None = None,
    ) -> list[IngestResult]:
        """Ingest all matching files from a directory."""
        dir_path = Path(dir_path)

        if not dir_path.is_dir():
            raise NotADirectoryError(f"Not a directory: {dir_path}")

        patterns = file_patterns or ["*.pdf", "*.doc", "*.docx", "*.txt", "*.jpg", "*.png"]

        files = []
        for pattern in patterns:
            if recursive:
                files.extend(dir_path.rglob(pattern))
            else:
                files.extend(dir_path.glob(pattern))

        results = []
        for file_path in files:
            try:
                result = await self.ingest_file(
                    file_path=file_path,
                    request_id=request_id,
                )
                results.append(result)
            except Exception as e:
                # Log error but continue with other files
                print(f"Error ingesting {file_path}: {e}")

        return results

    async def _calculate_checksum(self, file_path: Path) -> str:
        """Calculate SHA256 checksum of a file."""

        def _hash():
            sha256 = hashlib.sha256()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(8192), b""):
                    sha256.update(chunk)
            return sha256.hexdigest()

        return await asyncio.to_thread(_hash)

    async def _extract_text(self, file_path: Path, mime_type: str) -> str | None:
        """Auto-extract text from a file based on its type.

        - PDFs: try pdf-extract binary first (fast, lossless), fall back to OCR
        - DOCX: extract paragraphs + table text via python-docx
        - Images (JPG, PNG, TIFF): OCR via tesseract if available
        - Plain text: read directly
        """
        try:
            if mime_type == "application/pdf":
                return await self._extract_text_pdf(file_path)
            elif mime_type in (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ):
                return await asyncio.to_thread(self._extract_text_docx, file_path)
            elif mime_type and mime_type.startswith("image/"):
                return await self._extract_text_image(file_path)
            elif mime_type and mime_type.startswith("text/"):
                return await asyncio.to_thread(file_path.read_text, errors="replace")
            elif mime_type == "application/vnd.ms-outlook" or file_path.suffix.lower() == ".msg":
                return await asyncio.to_thread(self._extract_text_msg, file_path)
            return None
        except Exception:
            return None

    async def _extract_text_pdf(self, pdf_path: Path) -> str | None:
        """Extract text from PDF: try compiled binary first, fall back to OCR."""
        # Try the fast pdf-extract binary
        from .pdf_extract import extract_text

        result = await extract_text(pdf_path)
        if result and result.char_count > 50:
            # Clean up form feeds
            return "\n\n".join(p.strip() for p in result.text.split("\f") if p.strip())

        # Fall back to OCR for scanned PDFs
        try:
            from .ocr import OCREngine

            engine = OCREngine()
            ocr_result = await engine.process_pdf(pdf_path)
            if ocr_result.text and len(ocr_result.text.strip()) > 50:
                return ocr_result.text
        except (ImportError, Exception):
            pass

        return None

    def _extract_text_docx(self, docx_path: Path) -> str | None:
        """Extract text from DOCX including tables."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(str(docx_path))
            parts = []

            # Paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

            # Tables (FOIA privilege logs, fee estimates, etc.)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))

            return "\n".join(parts) if parts else None
        except Exception:
            return None

    async def _extract_text_image(self, image_path: Path) -> str | None:
        """Extract text from image via OCR (tesseract)."""
        try:
            import pytesseract
            from PIL import Image

            def _ocr():
                img = Image.open(image_path)
                return pytesseract.image_to_string(img)

            text = await asyncio.to_thread(_ocr)
            return text.strip() if text and len(text.strip()) > 10 else None
        except (ImportError, Exception):
            return None

    def _extract_text_msg(self, msg_path: Path) -> str | None:
        """Extract text from Outlook MSG files (email headers + body + attachment names)."""
        try:
            import extract_msg

            msg = extract_msg.Message(str(msg_path))
            parts = []

            # Email headers
            if msg.subject:
                parts.append(f"Subject: {msg.subject}")
            if msg.sender:
                parts.append(f"From: {msg.sender}")
            if msg.date:
                parts.append(f"Date: {msg.date}")
            if msg.to:
                parts.append(f"To: {msg.to}")

            parts.append("")

            # Body
            if msg.body:
                parts.append(msg.body)

            # Attachment names
            if msg.attachments:
                parts.append(f"\nAttachments ({len(msg.attachments)}):")
                for att in msg.attachments:
                    name = att.longFilename or att.shortFilename or "unnamed"
                    size = len(att.data) if att.data else 0
                    parts.append(f"  {name} ({size} bytes)")

            msg.close()
            text = "\n".join(parts)
            return text if len(text.strip()) > 10 else None
        except (ImportError, Exception):
            return None

    async def _get_pdf_page_count(self, pdf_path: Path) -> int:
        """Get the number of pages in a PDF."""

        def _count():
            from pypdf import PdfReader

            reader = PdfReader(str(pdf_path))
            return len(reader.pages)

        try:
            return await asyncio.to_thread(_count)
        except Exception:
            return 0
