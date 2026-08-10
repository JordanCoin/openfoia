"""Security regression tests: metadata stripping honesty.

The tool tells the user metadata was stripped. Anything it claims to remove
must actually be gone from the file — an audit dict that reports success for
a field it never touched is worse than not offering the feature.
"""

from __future__ import annotations

import zipfile

import pytest

# ---------------------------------------------------------------------------
# DOCX: company/manager live in docProps/app.xml, which python-docx cannot see
# ---------------------------------------------------------------------------


def _docx_with_extended_props(path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Body text")
    doc.core_properties.author = "Jane Source"
    doc.core_properties.last_modified_by = "Jane Source"
    doc.save(str(path))

    # python-docx does not write app.xml properties; inject them like Word does.
    app_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/'
        'extended-properties">'
        "<Company>Secret Ministry</Company>"
        "<Manager>Director Smith</Manager>"
        "<Application>Microsoft Word</Application>"
        "</Properties>"
    )
    _add_zip_member(path, "docProps/app.xml", app_xml)
    return path


def _add_zip_member(path, name, content):
    import os
    import shutil
    import tempfile

    fd, tmp = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    with zipfile.ZipFile(path) as src, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            if item.filename == name:
                continue
            dst.writestr(item, src.read(item.filename))
        dst.writestr(name, content)
    shutil.move(tmp, path)


def _read_zip_member(path, name):
    with zipfile.ZipFile(path) as z:
        if name not in z.namelist():
            return None
        return z.read(name).decode("utf-8", "replace")


def test_docx_strip_removes_company_and_manager(tmp_path):
    """_DOCX_SENSITIVE_ATTRS claimed these but never removed them."""
    from openfoia.pipeline.metadata import strip_metadata

    path = _docx_with_extended_props(tmp_path / "leak.docx")
    assert "Secret Ministry" in _read_zip_member(path, "docProps/app.xml")

    strip_metadata(path)

    app_xml = _read_zip_member(path, "docProps/app.xml") or ""
    assert "Secret Ministry" not in app_xml
    assert "Director Smith" not in app_xml


def test_docx_strip_removes_core_author(tmp_path):
    from openfoia.pipeline.metadata import strip_metadata

    path = _docx_with_extended_props(tmp_path / "leak.docx")
    strip_metadata(path)

    core = _read_zip_member(path, "docProps/core.xml") or ""
    assert "Jane Source" not in core


def test_docx_strip_keeps_document_readable(tmp_path):
    from docx import Document

    from openfoia.pipeline.metadata import strip_metadata

    path = _docx_with_extended_props(tmp_path / "leak.docx")
    strip_metadata(path)

    doc = Document(str(path))
    assert any("Body text" in p.text for p in doc.paragraphs)


def test_docx_report_does_not_claim_unremoved_fields(tmp_path):
    """Whatever `stripped` lists must genuinely be gone from the file."""
    from openfoia.pipeline.metadata import strip_metadata

    path = _docx_with_extended_props(tmp_path / "leak.docx")
    result = strip_metadata(path)

    blob = "".join(
        _read_zip_member(path, n) or ""
        for n in ("docProps/core.xml", "docProps/app.xml", "docProps/custom.xml")
    )
    if "company" in result.get("stripped", []):
        assert "Secret Ministry" not in blob
    if "manager" in result.get("stripped", []):
        assert "Director Smith" not in blob


# ---------------------------------------------------------------------------
# PDF: the XMP stream duplicates Author/Creator and survived /Info stripping
# ---------------------------------------------------------------------------


def _pdf_with_xmp(path):
    from pypdf import PdfWriter
    from pypdf.generic import ByteStringObject, DecodedStreamObject, NameObject

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.add_metadata({"/Author": "Jane Source", "/Producer": "SecretTool"})

    xmp = (
        b'<?xpacket begin="" id="W5M0MpCehiHzreSzNTczkc9d"?>'
        b'<x:xmpmeta xmlns:x="adobe:ns:meta/">'
        b'<rdf:RDF xmlns:rdf="http://www.w3.org/1999/02/22-rdf-syntax-ns#">'
        b'<rdf:Description xmlns:dc="http://purl.org/dc/elements/1.1/">'
        b"<dc:creator>Jane Source</dc:creator>"
        b"</rdf:Description></rdf:RDF></x:xmpmeta><?xpacket end='w'?>"
    )

    stream = DecodedStreamObject()
    stream.set_data(xmp)
    stream[NameObject("/Type")] = NameObject("/Metadata")
    stream[NameObject("/Subtype")] = NameObject("/XML")
    ref = writer._add_object(stream)
    writer._root_object[NameObject("/Metadata")] = ref
    # Keep a marker that is trivially greppable in the raw bytes.
    writer._info.get_object()[NameObject("/Title")] = ByteStringObject(b"Jane Source")

    with open(path, "wb") as f:
        writer.write(f)
    return path


def test_pdf_strip_removes_xmp_metadata(tmp_path):
    """XMP duplicates Author/Creator and was never cleared."""
    from openfoia.pipeline.metadata import strip_metadata

    path = _pdf_with_xmp(tmp_path / "leak.pdf")
    assert b"Jane Source" in path.read_bytes()

    strip_metadata(path)

    raw = path.read_bytes()
    assert b"dc:creator" not in raw
    assert b"Jane Source" not in raw


def test_pdf_strip_keeps_pages(tmp_path):
    from pypdf import PdfReader

    from openfoia.pipeline.metadata import strip_metadata

    path = _pdf_with_xmp(tmp_path / "leak.pdf")
    strip_metadata(path)

    assert len(PdfReader(str(path)).pages) == 1


def test_metadata_module_does_not_advertise_unstrippable_fields():
    """Every attr we list must be one we can actually clear."""
    from openfoia.pipeline import metadata

    # company/manager are only reachable via app.xml — the module must have a
    # dedicated code path for them, not just list them as core properties.
    assert hasattr(metadata, "_strip_docx_extended_props"), (
        "company/manager are listed as strippable but no app.xml handler exists"
    )


# ---------------------------------------------------------------------------
# OCR temp files must not escape `purge`
# ---------------------------------------------------------------------------


def test_ocr_temp_dir_is_inside_data_dir(tmp_path, monkeypatch):
    """pdf2image renders full-fidelity page images; /tmp escapes purge."""
    monkeypatch.setenv("OPENFOIA_DATA_DIR", str(tmp_path))
    from openfoia.db import get_data_dir
    from openfoia.pipeline.ocr import get_ocr_temp_dir

    d = get_ocr_temp_dir()
    assert get_data_dir() in d.parents or d.parent == get_data_dir()


@pytest.mark.skipif(__import__("os").name == "nt", reason="POSIX permissions")
def test_ocr_temp_dir_is_owner_only(tmp_path, monkeypatch):
    import os
    import stat

    monkeypatch.setenv("OPENFOIA_DATA_DIR", str(tmp_path))
    from openfoia.pipeline.ocr import get_ocr_temp_dir

    mode = stat.S_IMODE(get_ocr_temp_dir().stat().st_mode)
    assert mode & stat.S_IRWXG == 0
    assert mode & stat.S_IRWXO == 0
    assert os.name == "posix"
